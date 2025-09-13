"""
Procesador de documentos para extracción automática de Q&A
Soporta PDF, DOCX, TXT y otros formatos comunes
"""

import os
import re
from pathlib import Path
from typing import List, Dict, Any, Optional, Union
from abc import ABC, abstractmethod

import PyPDF2
from docx import Document
import nltk
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document as LangchainDocument

from config.settings import settings
from src.utils.logger import get_logger
from src.utils.data_models import QAItem, QABatch, GenerationRequest
from src.generators.prompt_generator import PromptQAGenerator

logger = get_logger(__name__)

# Descargar recursos de NLTK necesarios
try:
    nltk.download('punkt', quiet=True)
    nltk.download('stopwords', quiet=True)
except:
    logger.warning("No se pudieron descargar recursos de NLTK")

class DocumentReader(ABC):
    """Clase base para lectores de documentos"""
    
    @abstractmethod
    def read(self, file_path: str) -> str:
        """Leer documento y retornar texto"""
        pass
    
    @abstractmethod
    def get_supported_extensions(self) -> List[str]:
        """Obtener extensiones soportadas"""
        pass

class PDFReader(DocumentReader):
    """Lector de archivos PDF"""
    
    def read(self, file_path: str) -> str:
        """Leer archivo PDF"""
        try:
            text = ""
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
            return text
        except Exception as e:
            logger.error(f"Error leyendo PDF {file_path}: {e}")
            raise
    
    def get_supported_extensions(self) -> List[str]:
        return ['.pdf']

class DOCXReader(DocumentReader):
    """Lector de archivos DOCX"""
    
    def read(self, file_path: str) -> str:
        """Leer archivo DOCX"""
        try:
            doc = Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        except Exception as e:
            logger.error(f"Error leyendo DOCX {file_path}: {e}")
            raise
    
    def get_supported_extensions(self) -> List[str]:
        return ['.docx', '.doc']

class TXTReader(DocumentReader):
    """Lector de archivos de texto plano"""
    
    def read(self, file_path: str) -> str:
        """Leer archivo de texto"""
        try:
            encodings = ['utf-8', 'latin-1', 'cp1252']
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        return file.read()
                except UnicodeDecodeError:
                    continue
            raise ValueError(f"No se pudo decodificar el archivo {file_path}")
        except Exception as e:
            logger.error(f"Error leyendo TXT {file_path}: {e}")
            raise
    
    def get_supported_extensions(self) -> List[str]:
        return ['.txt', '.md', '.rtf']

class DocumentChunker:
    """Segmentador de texto en chunks para procesamiento"""
    
    def __init__(self, chunk_size: int = None, chunk_overlap: int = None):
        self.chunk_size = chunk_size or settings.CHUNK_SIZE
        self.chunk_overlap = chunk_overlap or settings.CHUNK_OVERLAP
        
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=self.chunk_size,
            chunk_overlap=self.chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", ".", "!", "?", ";", ":", " ", ""]
        )
    
    def chunk_text(self, text: str, metadata: Dict[str, Any] = None) -> List[LangchainDocument]:
        """Dividir texto en chunks"""
        try:
            chunks = self.text_splitter.split_text(text)
            documents = []
            
            for i, chunk in enumerate(chunks):
                doc_metadata = {
                    "chunk_index": i,
                    "chunk_size": len(chunk),
                    **(metadata or {})
                }
                documents.append(LangchainDocument(page_content=chunk, metadata=doc_metadata))
            
            logger.info(f"Texto dividido en {len(documents)} chunks")
            return documents
        except Exception as e:
            logger.error(f"Error dividiendo texto en chunks: {e}")
            raise

class QAExtractor:
    """Extractor de preguntas y respuestas de texto"""
    
    def __init__(self):
        self.prompt_generator = PromptQAGenerator()
    
    def extract_existing_qa(self, text: str) -> List[Dict[str, str]]:
        """Extraer Q&A que ya existen en el texto (si los hay)"""
        qa_pairs = []
        
        # Patrones para identificar Q&A existentes
        patterns = [
            r'P(?:regunta)?[:\-\s]*(.+?)\s*R(?:espuesta)?[:\-\s]*(.+?)(?=P(?:regunta)?[:\-\s]|$)',
            r'Q[:\-\s]*(.+?)\s*A[:\-\s]*(.+?)(?=Q[:\-\s]|$)',
            r'(\d+\.\s*.+?\?)\s*(.+?)(?=\d+\.|$)',
            r'¿(.+?\?)\s*(.+?)(?=¿|$)'
        ]
        
        for pattern in patterns:
            matches = re.findall(pattern, text, re.DOTALL | re.IGNORECASE)
            for question, answer in matches:
                qa_pairs.append({
                    "pregunta": question.strip(),
                    "respuesta": answer.strip()
                })
        
        # Filtrar duplicados y validar calidad
        filtered_pairs = []
        seen_questions = set()
        
        for qa in qa_pairs:
            question = qa["pregunta"].lower()
            if (question not in seen_questions and 
                len(qa["pregunta"]) > 10 and 
                len(qa["respuesta"]) > 20 and
                "?" in qa["pregunta"]):
                filtered_pairs.append(qa)
                seen_questions.add(question)
        
        return filtered_pairs
    
    async def generate_qa_from_chunk(
        self,
        chunk: LangchainDocument,
        categoria: str = "general",
        nivel: str = "intermedio",
        num_preguntas: int = 5,
        modelo: str = "openai"
    ) -> List[QAItem]:
        """Generar Q&A a partir de un chunk de texto"""
        
        # Crear prompt específico para extracción de documentos
        prompt = f"""
        Basándote en el siguiente texto, genera {num_preguntas} preguntas y respuestas específicas y relevantes.
        
        TEXTO:
        {chunk.page_content}
        
        Instrucciones:
        - Las preguntas deben ser respondibles directamente con la información del texto
        - Las respuestas deben ser precisas y basadas únicamente en el contenido proporcionado
        - Incluye preguntas de diferentes tipos: factual, conceptual, y de comprensión
        - Usa un nivel de dificultad: {nivel}
        - Categoría: {categoria}
        
        Formato de respuesta:
        PREGUNTA: [pregunta aquí]
        RESPUESTA: [respuesta aquí]
        ---
        """
        
        # Crear request para el generador
        request = GenerationRequest(
            tipo="documento",
            prompt=prompt,
            categoria=categoria,
            nivel=nivel,
            num_preguntas=num_preguntas,
            modelo=modelo
        )
        
        # Generar Q&A
        provider = self.prompt_generator.get_provider(modelo)
        response = await provider.generate_qa(prompt)
        
        # Parsear respuesta
        qa_items = self.prompt_generator.parse_qa_response(response, request)
        
        # Agregar metadatos del chunk
        for item in qa_items:
            item.metadatos.update({
                "fuente_chunk": chunk.metadata.get("chunk_index", 0),
                "documento_fuente": chunk.metadata.get("source", ""),
                "metodo_extraccion": "generacion_automatica"
            })
        
        return qa_items

class DocumentQAProcessor:
    """Procesador principal de documentos para generación de Q&A"""
    
    def __init__(self):
        self.readers = {
            '.pdf': PDFReader(),
            '.docx': DOCXReader(),
            '.doc': DOCXReader(),
            '.txt': TXTReader(),
            '.md': TXTReader(),
            '.rtf': TXTReader()
        }
        self.chunker = DocumentChunker()
        self.qa_extractor = QAExtractor()
    
    def get_supported_formats(self) -> List[str]:
        """Obtener formatos de archivo soportados"""
        return list(self.readers.keys())
    
    def validate_document(self, file_path: str) -> bool:
        """Validar que el documento sea procesable"""
        try:
            path = Path(file_path)
            
            # Verificar que existe
            if not path.exists():
                logger.error(f"Archivo no existe: {file_path}")
                return False
            
            # Verificar extensión
            if path.suffix.lower() not in self.readers:
                logger.error(f"Formato no soportado: {path.suffix}")
                return False
            
            # Verificar tamaño
            size_mb = path.stat().st_size / (1024 * 1024)
            if size_mb > settings.MAX_DOCUMENT_SIZE_MB:
                logger.error(f"Archivo muy grande: {size_mb:.2f}MB > {settings.MAX_DOCUMENT_SIZE_MB}MB")
                return False
            
            return True
        except Exception as e:
            logger.error(f"Error validando documento: {e}")
            return False
    
    def read_document(self, file_path: str) -> str:
        """Leer documento usando el reader apropiado"""
        path = Path(file_path)
        extension = path.suffix.lower()
        
        if extension not in self.readers:
            raise ValueError(f"Formato no soportado: {extension}")
        
        reader = self.readers[extension]
        return reader.read(file_path)
    
    async def process_document(
        self,
        file_path: str,
        categoria: str = "general",
        nivel: str = "intermedio",
        preguntas_por_chunk: int = 3,
        modelo: str = "openai",
        extraer_existente: bool = True
    ) -> QABatch:
        """Procesar documento completo y generar Q&A"""
        
        logger.info(f"Procesando documento: {file_path}")
        
        # Validar documento
        if not self.validate_document(file_path):
            raise ValueError(f"Documento inválido: {file_path}")
        
        # Leer contenido
        text = self.read_document(file_path)
        logger.info(f"Texto extraído: {len(text)} caracteres")
        
        # Crear metadatos del documento
        doc_metadata = {
            "source": file_path,
            "filename": Path(file_path).name,
            "extension": Path(file_path).suffix,
            "size_chars": len(text)
        }
        
        # Dividir en chunks
        chunks = self.chunker.chunk_text(text, doc_metadata)
        
        # Extraer Q&A existentes si se solicita
        existing_qa = []
        if extraer_existente:
            logger.info("Extrayendo Q&A existentes del documento")
            existing_pairs = self.qa_extractor.extract_existing_qa(text)
            for pair in existing_pairs:
                qa_item = QAItem(
                    pregunta=pair["pregunta"],
                    respuesta=pair["respuesta"],
                    categoria=categoria,
                    nivel=nivel,
                    fuentes=[file_path],
                    metadatos={
                        "metodo_extraccion": "extraccion_directa",
                        "documento_fuente": file_path
                    }
                )
                existing_qa.append(qa_item)
        
        # Generar Q&A de chunks
        generated_qa = []
        for i, chunk in enumerate(chunks):
            logger.info(f"Procesando chunk {i+1}/{len(chunks)}")
            try:
                chunk_qa = await self.qa_extractor.generate_qa_from_chunk(
                    chunk, categoria, nivel, preguntas_por_chunk, modelo
                )
                generated_qa.extend(chunk_qa)
            except Exception as e:
                logger.warning(f"Error procesando chunk {i+1}: {e}")
                continue
        
        # Combinar todos los Q&A
        all_qa = existing_qa + generated_qa
        
        # Crear batch
        batch = QABatch(
            items=all_qa,
            origen="documento",
            documento_fuente=file_path,
            parametros_generacion={
                "categoria": categoria,
                "nivel": nivel,
                "preguntas_por_chunk": preguntas_por_chunk,
                "modelo": modelo,
                "extraer_existente": extraer_existente,
                "num_chunks": len(chunks)
            }
        )
        
        batch.estadisticas = batch.get_stats()
        
        logger.info(f"Procesamiento completado: {len(all_qa)} Q&A generados")
        logger.info(f"  - Extraídos existentes: {len(existing_qa)}")
        logger.info(f"  - Generados automáticamente: {len(generated_qa)}")
        
        return batch

# Función de conveniencia
async def process_document_to_qa(
    file_path: str,
    categoria: str = "general",
    nivel: str = "intermedio",
    preguntas_por_chunk: int = 3,
    modelo: str = "openai"
) -> QABatch:
    """Función de conveniencia para procesar documento"""
    
    processor = DocumentQAProcessor()
    return await processor.process_document(
        file_path, categoria, nivel, preguntas_por_chunk, modelo
    )